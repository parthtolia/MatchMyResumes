"""
Generate 20 ATS-friendly resume templates as properly formatted DOCX files.
Each template uses clean formatting that ATS parsers can read:
- Standard section headings (Experience, Education, Skills, etc.)
- No tables, text boxes, headers/footers, or images
- Single-column layout
- Standard fonts (Calibri, Arial, Times New Roman)
- Proper heading styles for section parsing
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "public", "templates")
os.makedirs(OUT_DIR, exist_ok=True)


def set_margins(doc, top=0.7, bottom=0.7, left=0.75, right=0.75):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_name(doc, name, font_name="Calibri", size=22, color=None, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.space_after = Pt(2)
    p.space_before = Pt(0)
    run = p.add_run(name)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_contact(doc, contact_text, font_name="Calibri", size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.space_after = Pt(6)
    p.space_before = Pt(0)
    run = p.add_run(contact_text)
    run.font.name = font_name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_section_heading(doc, title, font_name="Calibri", size=12, color=None, bold=True, uppercase=True, underline=False, border_bottom=False):
    text = title.upper() if uppercase else title
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)

    if border_bottom:
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '1',
            qn('w:color'): '333333'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    return p


def add_entry(doc, title, subtitle="", date="", font_name="Calibri", title_size=11, sub_size=10):
    # Title + date on same line
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(1)
    run = p.add_run(title)
    run.font.name = font_name
    run.font.size = Pt(title_size)
    run.font.bold = True
    if date:
        run = p.add_run(f"  |  {date}")
        run.font.name = font_name
        run.font.size = Pt(sub_size)
        run.font.color.rgb = RGBColor(100, 100, 100)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        run = p2.add_run(subtitle)
        run.font.name = font_name
        run.font.size = Pt(sub_size)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)


def add_bullet(doc, text, font_name="Calibri", size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    p.clear()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)


def add_body(doc, text, font_name="Calibri", size=10, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_skills_line(doc, label, skills, font_name="Calibri", size=10):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(f"{label}: ")
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = True
    run = p.add_run(skills)
    run.font.name = font_name
    run.font.size = Pt(size)


def add_separator(doc):
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# =============================================================================
# TEMPLATE DEFINITIONS
# =============================================================================

TEMPLATES = []

# ---------------------------------------------------------------------------
# 1. Modern Clean
# ---------------------------------------------------------------------------
def gen_modern_clean():
    doc = Document()
    set_margins(doc)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 24, (44, 62, 80))
    add_contact(doc, "[City, State] | [Phone] | [Email] | [LinkedIn URL]", "Calibri", 10, color=(100, 100, 100))
    add_separator(doc)

    add_section_heading(doc, "Professional Summary", "Calibri", 12, (44, 62, 80), border_bottom=True)
    add_body(doc, "Results-driven professional with [X] years of experience in [industry/field]. Proven track record of [key achievement]. Seeking to leverage expertise in [skill area] to drive impact at [target company type].")

    add_section_heading(doc, "Professional Experience", "Calibri", 12, (44, 62, 80), border_bottom=True)
    add_entry(doc, "[Job Title]", "[Company Name] — [City, State]", "[Start Date] – [End Date]")
    add_bullet(doc, "Achieved [specific result] by implementing [action], resulting in [quantified outcome, e.g., 25% increase in revenue]")
    add_bullet(doc, "Led cross-functional team of [X] members to deliver [project/initiative] [on time/under budget]")
    add_bullet(doc, "Streamlined [process] using [tool/methodology], reducing [metric] by [X]%")

    add_entry(doc, "[Job Title]", "[Company Name] — [City, State]", "[Start Date] – [End Date]")
    add_bullet(doc, "Managed [responsibility area] supporting [X] clients/users across [scope]")
    add_bullet(doc, "Developed and executed [strategy/plan] that improved [KPI] by [X]%")
    add_bullet(doc, "Collaborated with [stakeholders] to launch [product/feature/campaign]")

    add_section_heading(doc, "Education", "Calibri", 12, (44, 62, 80), border_bottom=True)
    add_entry(doc, "[Degree] in [Major]", "[University Name] — [City, State]", "[Graduation Year]")
    add_body(doc, "GPA: [X.X/4.0] | Relevant Coursework: [Course 1], [Course 2], [Course 3]")

    add_section_heading(doc, "Skills", "Calibri", 12, (44, 62, 80), border_bottom=True)
    add_skills_line(doc, "Technical", "[Skill 1], [Skill 2], [Skill 3], [Skill 4], [Skill 5]")
    add_skills_line(doc, "Tools", "[Tool 1], [Tool 2], [Tool 3], [Tool 4]")
    add_skills_line(doc, "Soft Skills", "[Leadership], [Communication], [Problem-Solving], [Teamwork]")

    add_section_heading(doc, "Certifications", "Calibri", 12, (44, 62, 80), border_bottom=True)
    add_body(doc, "[Certification Name] — [Issuing Organization] — [Year]")
    add_body(doc, "[Certification Name] — [Issuing Organization] — [Year]")

    return doc

TEMPLATES.append(("modern-clean-ats-resume.docx", gen_modern_clean))


# ---------------------------------------------------------------------------
# 2. Modern Minimal
# ---------------------------------------------------------------------------
def gen_modern_minimal():
    doc = Document()
    set_margins(doc)
    add_name(doc, "[YOUR FULL NAME]", "Arial", 22, (0, 128, 128))
    add_contact(doc, "[Email] • [Phone] • [LinkedIn] • [City, State]", "Arial", 10, color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Summary", "Arial", 11, (0, 128, 128), border_bottom=False, underline=True)
    add_body(doc, "[X]-year [industry] professional specializing in [core competency]. Track record of [key achievement]. Passionate about [area of interest] with strong expertise in [relevant skill].", "Arial")

    add_section_heading(doc, "Experience", "Arial", 11, (0, 128, 128), underline=True)
    add_entry(doc, "[Job Title]", "[Company Name], [City, State]", "[Month Year] – Present", "Arial")
    add_bullet(doc, "Drove [X]% improvement in [metric] through [specific initiative]", "Arial")
    add_bullet(doc, "Built and maintained [system/process] serving [X] users/clients", "Arial")
    add_bullet(doc, "Partnered with [department] to implement [solution] reducing costs by $[X]K", "Arial")

    add_entry(doc, "[Job Title]", "[Company Name], [City, State]", "[Month Year] – [Month Year]", "Arial")
    add_bullet(doc, "Spearheaded [project] from concept to delivery in [X] months", "Arial")
    add_bullet(doc, "Optimized [workflow] resulting in [X] hours saved per week", "Arial")

    add_section_heading(doc, "Education", "Arial", 11, (0, 128, 128), underline=True)
    add_entry(doc, "[Degree] in [Field]", "[University Name]", "[Year]", "Arial")

    add_section_heading(doc, "Skills", "Arial", 11, (0, 128, 128), underline=True)
    add_body(doc, "[Skill 1] • [Skill 2] • [Skill 3] • [Skill 4] • [Skill 5] • [Skill 6] • [Skill 7] • [Skill 8]", "Arial")

    return doc

TEMPLATES.append(("modern-minimal-ats-resume.docx", gen_modern_minimal))


# ---------------------------------------------------------------------------
# 3. Simple Classic
# ---------------------------------------------------------------------------
def gen_simple_classic():
    doc = Document()
    set_margins(doc, 1, 1, 1, 1)
    add_name(doc, "[YOUR FULL NAME]", "Times New Roman", 20, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_contact(doc, "[Address] | [Phone] | [Email]", "Times New Roman", 11, WD_ALIGN_PARAGRAPH.CENTER)
    add_separator(doc)

    add_section_heading(doc, "Objective", "Times New Roman", 12, bold=True, uppercase=True, border_bottom=True)
    add_body(doc, "Dedicated [job title/field] professional with [X] years of experience seeking a [target position] role where I can apply my expertise in [skill 1], [skill 2], and [skill 3] to contribute to [company goal/mission].", "Times New Roman", 11)

    add_section_heading(doc, "Work Experience", "Times New Roman", 12, bold=True, uppercase=True, border_bottom=True)
    add_entry(doc, "[Job Title]", "[Company Name], [City, State]", "[Start] – [End]", "Times New Roman", 11, 10)
    add_bullet(doc, "Managed [responsibility] for [scope], achieving [measurable result]", "Times New Roman", 11)
    add_bullet(doc, "Trained and mentored [X] new team members on [process/tool]", "Times New Roman", 11)
    add_bullet(doc, "Reduced [problem metric] by [X]% through [action taken]", "Times New Roman", 11)

    add_entry(doc, "[Job Title]", "[Company Name], [City, State]", "[Start] – [End]", "Times New Roman", 11, 10)
    add_bullet(doc, "Coordinated [activity] across [X] departments ensuring [outcome]", "Times New Roman", 11)
    add_bullet(doc, "Processed [X]+ [items/requests] per [time period] with [X]% accuracy", "Times New Roman", 11)

    add_section_heading(doc, "Education", "Times New Roman", 12, bold=True, uppercase=True, border_bottom=True)
    add_entry(doc, "[Degree] in [Major]", "[University Name], [City, State]", "[Graduation Year]", "Times New Roman", 11, 10)

    add_section_heading(doc, "Skills", "Times New Roman", 12, bold=True, uppercase=True, border_bottom=True)
    add_body(doc, "[Skill 1], [Skill 2], [Skill 3], [Skill 4], [Skill 5], [Skill 6], [Skill 7], [Skill 8]", "Times New Roman", 11)

    add_section_heading(doc, "References", "Times New Roman", 12, bold=True, uppercase=True, border_bottom=True)
    add_body(doc, "Available upon request.", "Times New Roman", 11)

    return doc

TEMPLATES.append(("simple-classic-ats-resume.docx", gen_simple_classic))


# ---------------------------------------------------------------------------
# 4. Simple One Column
# ---------------------------------------------------------------------------
def gen_simple_one_column():
    doc = Document()
    set_margins(doc)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 20)
    add_contact(doc, "[Phone] | [Email] | [LinkedIn URL] | [City, State]")
    add_separator(doc)

    add_section_heading(doc, "Professional Summary", border_bottom=True)
    add_body(doc, "Accomplished [profession] with [X]+ years in [industry]. Expert in [area 1] and [area 2]. Known for [distinguishing quality] and delivering [type of results].")

    add_section_heading(doc, "Work Experience", border_bottom=True)
    add_entry(doc, "[Job Title]", "[Company] — [Location]", "[Date Range]")
    add_bullet(doc, "Increased [metric] by [X]% by designing and implementing [initiative]")
    add_bullet(doc, "Managed annual budget of $[X]M for [department/project]")
    add_bullet(doc, "Delivered [project/product] [ahead of schedule/under budget], impacting [X] users")

    add_entry(doc, "[Job Title]", "[Company] — [Location]", "[Date Range]")
    add_bullet(doc, "Automated [process] using [tool], saving [X] hours/month")
    add_bullet(doc, "Received [award/recognition] for [achievement]")

    add_section_heading(doc, "Education", border_bottom=True)
    add_entry(doc, "[Degree], [Major]", "[University]", "[Year]")

    add_section_heading(doc, "Technical Skills", border_bottom=True)
    add_body(doc, "[Skill 1], [Skill 2], [Skill 3], [Skill 4], [Skill 5], [Skill 6]")

    add_section_heading(doc, "Certifications", border_bottom=True)
    add_body(doc, "[Certification] — [Organization] — [Year]")

    return doc

TEMPLATES.append(("simple-one-column-ats-resume.docx", gen_simple_one_column))


# ---------------------------------------------------------------------------
# 5. Professional Standard
# ---------------------------------------------------------------------------
def gen_professional_standard():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (0, 51, 102))
    add_contact(doc, "[Job Title] | [Phone] | [Email] | [LinkedIn] | [City, State]", size=10, color=(80, 80, 80))
    add_separator(doc)

    add_section_heading(doc, "Professional Profile", color=(0, 51, 102), border_bottom=True)
    add_body(doc, "Senior [profession] with [X]+ years of progressive experience in [industry]. Proven ability to [core competency 1] and [core competency 2]. Recognized for [distinguishing quality]. Seeking to contribute to [target role/company type].")

    add_section_heading(doc, "Core Competencies", color=(0, 51, 102), border_bottom=True)
    add_body(doc, "[Competency 1] • [Competency 2] • [Competency 3] • [Competency 4]\n[Competency 5] • [Competency 6] • [Competency 7] • [Competency 8]")

    add_section_heading(doc, "Professional Experience", color=(0, 51, 102), border_bottom=True)
    add_entry(doc, "[Job Title]", "[Company Name] — [City, State]", "[Start] – [End]")
    add_bullet(doc, "Led [initiative] that generated $[X]M in [revenue/savings] over [time period]")
    add_bullet(doc, "Directed team of [X] professionals across [X] locations/departments")
    add_bullet(doc, "Established [program/system] that improved [metric] by [X]%")
    add_bullet(doc, "Presented [findings/strategy] to [C-suite/board/stakeholders]")

    add_entry(doc, "[Job Title]", "[Company Name] — [City, State]", "[Start] – [End]")
    add_bullet(doc, "Developed [solution] addressing [business challenge], resulting in [outcome]")
    add_bullet(doc, "Negotiated [contracts/partnerships] worth $[X]M with [clients/vendors]")
    add_bullet(doc, "Implemented [methodology/framework] across [scope], increasing [KPI] by [X]%")

    add_section_heading(doc, "Education", color=(0, 51, 102), border_bottom=True)
    add_entry(doc, "[Degree] in [Major]", "[University Name]", "[Year]")
    add_entry(doc, "[Degree] in [Major]", "[University Name]", "[Year]")

    add_section_heading(doc, "Certifications & Awards", color=(0, 51, 102), border_bottom=True)
    add_body(doc, "[Certification/Award] — [Organization] — [Year]")
    add_body(doc, "[Certification/Award] — [Organization] — [Year]")

    return doc

TEMPLATES.append(("professional-standard-ats-resume.docx", gen_professional_standard))


# ---------------------------------------------------------------------------
# 6. Corporate Professional
# ---------------------------------------------------------------------------
def gen_corporate_professional():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 24, (30, 30, 30))
    add_contact(doc, "[Title] | [Email] | [Phone] | [LinkedIn]", size=10, color=(100, 100, 100))
    add_separator(doc)

    add_section_heading(doc, "Executive Summary", color=(30, 30, 30), border_bottom=True)
    add_body(doc, "Strategic [profession] with [X]+ years of leadership experience in [sector]. Track record of driving [growth/transformation/efficiency] across [scope]. Expert in [domain 1], [domain 2], and [domain 3].")

    add_section_heading(doc, "Key Achievements", color=(30, 30, 30), border_bottom=True)
    add_bullet(doc, "Grew [business unit/revenue] from $[X]M to $[X]M in [timeframe]")
    add_bullet(doc, "Led M&A integration of [X] companies, retaining [X]% of key talent")
    add_bullet(doc, "Launched [product/service] generating $[X]M in first-year revenue")

    add_section_heading(doc, "Professional Experience", color=(30, 30, 30), border_bottom=True)
    add_entry(doc, "[Senior Title]", "[Corporation] — [Location]", "[Start] – Present")
    add_bullet(doc, "Oversee [division/function] with P&L responsibility of $[X]M")
    add_bullet(doc, "Developed and executed [X]-year strategic plan achieving [X]% CAGR")
    add_bullet(doc, "Built and scaled team from [X] to [X] across [X] offices/countries")

    add_entry(doc, "[Title]", "[Corporation] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Restructured [department] reducing operational costs by $[X]M annually")
    add_bullet(doc, "Forged strategic partnerships with [X] enterprise clients")

    add_section_heading(doc, "Education", color=(30, 30, 30), border_bottom=True)
    add_entry(doc, "MBA, [Concentration]", "[Business School]", "[Year]")
    add_entry(doc, "[Bachelor's Degree] in [Major]", "[University]", "[Year]")

    add_section_heading(doc, "Board & Advisory Roles", color=(30, 30, 30), border_bottom=True)
    add_body(doc, "[Organization Name] — [Role] — [Year]–Present")

    return doc

TEMPLATES.append(("corporate-professional-ats-resume.docx", gen_corporate_professional))


# ---------------------------------------------------------------------------
# 7. Fresh Graduate
# ---------------------------------------------------------------------------
def gen_fresh_graduate():
    doc = Document()
    set_margins(doc)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (85, 85, 85))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [Portfolio URL (optional)]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Objective", color=(85, 85, 85), border_bottom=True)
    add_body(doc, "Motivated recent graduate with a [Degree] in [Major] from [University]. Eager to apply academic knowledge in [skill area 1] and [skill area 2] to a [target role] position. Strong foundation in [relevant coursework/projects].")

    add_section_heading(doc, "Education", color=(85, 85, 85), border_bottom=True)
    add_entry(doc, "[Degree] in [Major]", "[University Name] — [City, State]", "[Graduation Month Year]")
    add_body(doc, "GPA: [X.X/4.0] | Dean's List: [Semesters] | Honors: [Any honors]")
    add_body(doc, "Relevant Coursework: [Course 1], [Course 2], [Course 3], [Course 4]")

    add_section_heading(doc, "Projects", color=(85, 85, 85), border_bottom=True)
    add_entry(doc, "[Project Name]", "[Course/Personal/Hackathon]", "[Date]")
    add_bullet(doc, "Developed [what] using [technologies/tools] to [solve what problem]")
    add_bullet(doc, "Achieved [result, e.g., 95% accuracy, 2nd place, 500+ users]")

    add_entry(doc, "[Project Name]", "[Course/Personal/Hackathon]", "[Date]")
    add_bullet(doc, "Built [application/system] that [functionality] for [audience/use case]")
    add_bullet(doc, "Utilized [framework/language] with [methodology, e.g., Agile, TDD]")

    add_section_heading(doc, "Internship / Work Experience", color=(85, 85, 85), border_bottom=True)
    add_entry(doc, "[Intern / Part-Time Title]", "[Company] — [Location]", "[Date Range]")
    add_bullet(doc, "Assisted [team/department] with [task], contributing to [outcome]")
    add_bullet(doc, "Created [deliverable] used by [X] team members/stakeholders")

    add_section_heading(doc, "Skills", color=(85, 85, 85), border_bottom=True)
    add_skills_line(doc, "Technical", "[Language 1], [Language 2], [Framework 1], [Tool 1], [Tool 2]")
    add_skills_line(doc, "Soft Skills", "[Communication], [Teamwork], [Time Management], [Problem-Solving]")

    add_section_heading(doc, "Extracurricular Activities", color=(85, 85, 85), border_bottom=True)
    add_body(doc, "[Club/Organization] — [Role] — [Date Range]")
    add_body(doc, "[Volunteer Work] — [Organization] — [Date Range]")

    return doc

TEMPLATES.append(("fresh-graduate-ats-resume.docx", gen_fresh_graduate))


# ---------------------------------------------------------------------------
# 8. Internship Ready
# ---------------------------------------------------------------------------
def gen_internship_ready():
    doc = Document()
    set_margins(doc)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 20, (200, 150, 50))
    add_contact(doc, "[University Email] | [Phone] | [LinkedIn] | [GitHub (optional)]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Objective", color=(160, 120, 40), border_bottom=True)
    add_body(doc, "[Year] [Major] student at [University] seeking a [Summer/Fall] [Year] internship in [field]. Experienced with [tool/language 1] and [tool/language 2] through coursework and personal projects.")

    add_section_heading(doc, "Education", color=(160, 120, 40), border_bottom=True)
    add_entry(doc, "[Degree] in [Major] (Expected [Graduation Date])", "[University Name]", "GPA: [X.X]")
    add_body(doc, "Relevant Coursework: [Course 1], [Course 2], [Course 3], [Course 4], [Course 5]")

    add_section_heading(doc, "Projects", color=(160, 120, 40), border_bottom=True)
    add_entry(doc, "[Project Name]", "[Technologies Used]", "[Date]")
    add_bullet(doc, "[What you built] — [What it does] — [Impact/result]")
    add_bullet(doc, "[Technical detail about implementation]")

    add_entry(doc, "[Project Name]", "[Technologies Used]", "[Date]")
    add_bullet(doc, "[Description of project and your role]")

    add_section_heading(doc, "Experience", color=(160, 120, 40), border_bottom=True)
    add_entry(doc, "[Role, e.g., Teaching Assistant / Lab Assistant / Tutor]", "[Department/Organization]", "[Date Range]")
    add_bullet(doc, "[Responsibility] for [X] students/participants")
    add_bullet(doc, "[Achievement or contribution]")

    add_section_heading(doc, "Technical Skills", color=(160, 120, 40), border_bottom=True)
    add_skills_line(doc, "Languages", "[Python], [Java], [JavaScript], [C++], [SQL]")
    add_skills_line(doc, "Frameworks", "[React], [Node.js], [Django], [Flask]")
    add_skills_line(doc, "Tools", "[Git], [Docker], [AWS], [VS Code], [Jupyter]")

    add_section_heading(doc, "Leadership & Activities", color=(160, 120, 40), border_bottom=True)
    add_body(doc, "[Organization] — [Position] — [Date Range]")
    add_body(doc, "[Hackathon/Competition] — [Result/Award] — [Date]")

    return doc

TEMPLATES.append(("internship-ready-ats-resume.docx", gen_internship_ready))


# ---------------------------------------------------------------------------
# 9. Software Developer
# ---------------------------------------------------------------------------
def gen_software_developer():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (108, 92, 231))
    add_contact(doc, "[Email] | [Phone] | [GitHub] | [LinkedIn] | [Portfolio]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Summary", color=(108, 92, 231), border_bottom=True)
    add_body(doc, "Software Engineer with [X]+ years of experience building scalable [web/mobile/cloud] applications. Proficient in [primary language] and [secondary language]. Passionate about [area, e.g., distributed systems, developer tooling, open source].")

    add_section_heading(doc, "Technical Skills", color=(108, 92, 231), border_bottom=True)
    add_skills_line(doc, "Languages", "[Python], [TypeScript], [Go], [Java], [SQL]")
    add_skills_line(doc, "Frontend", "[React], [Next.js], [Tailwind CSS], [HTML5/CSS3]")
    add_skills_line(doc, "Backend", "[Node.js], [FastAPI], [Django], [Express], [GraphQL]")
    add_skills_line(doc, "Databases", "[PostgreSQL], [MongoDB], [Redis], [DynamoDB]")
    add_skills_line(doc, "DevOps & Cloud", "[AWS], [Docker], [Kubernetes], [CI/CD], [Terraform]")
    add_skills_line(doc, "Tools", "[Git], [Jira], [Datadog], [Figma], [Postman]")

    add_section_heading(doc, "Professional Experience", color=(108, 92, 231), border_bottom=True)
    add_entry(doc, "[Software Engineer / Senior SWE]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Architected and shipped [feature/service] handling [X]K+ requests/second")
    add_bullet(doc, "Reduced API latency by [X]% through [caching/query optimization/refactoring]")
    add_bullet(doc, "Implemented CI/CD pipeline reducing deployment time from [X] hours to [X] minutes")
    add_bullet(doc, "Mentored [X] junior engineers through code reviews and pair programming")

    add_entry(doc, "[Software Engineer]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Developed [microservice/module] in [language] processing [X]M+ records daily")
    add_bullet(doc, "Migrated legacy [system] to [modern stack], improving performance by [X]%")
    add_bullet(doc, "Contributed to open-source [project/library] with [X]+ GitHub stars")

    add_section_heading(doc, "Projects", color=(108, 92, 231), border_bottom=True)
    add_entry(doc, "[Project Name]", "[Tech Stack]", "[Link/GitHub]")
    add_bullet(doc, "[Description] — [X] users/stars — [Key technical highlight]")

    add_section_heading(doc, "Education", color=(108, 92, 231), border_bottom=True)
    add_entry(doc, "[BS/MS] in [Computer Science / Software Engineering]", "[University]", "[Year]")

    return doc

TEMPLATES.append(("software-developer-ats-resume.docx", gen_software_developer))


# ---------------------------------------------------------------------------
# 10. Data Science
# ---------------------------------------------------------------------------
def gen_data_science():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (0, 128, 128))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [GitHub] | [Kaggle (optional)]", color=(100, 100, 100))
    add_separator(doc)

    add_section_heading(doc, "Summary", color=(0, 128, 128), border_bottom=True)
    add_body(doc, "Data Scientist with [X]+ years of experience applying statistical modeling and machine learning to solve [business domain] problems. Published [X] papers in [venues]. Skilled in translating complex data insights into actionable business recommendations.")

    add_section_heading(doc, "Technical Skills", color=(0, 128, 128), border_bottom=True)
    add_skills_line(doc, "Languages", "Python, R, SQL, Scala")
    add_skills_line(doc, "ML/AI", "Scikit-learn, TensorFlow, PyTorch, XGBoost, Hugging Face")
    add_skills_line(doc, "Data Engineering", "Spark, Airflow, dbt, Kafka, ETL pipelines")
    add_skills_line(doc, "Visualization", "Matplotlib, Seaborn, Plotly, Tableau, Looker")
    add_skills_line(doc, "Cloud & Tools", "AWS SageMaker, GCP Vertex AI, MLflow, Docker, Git")
    add_skills_line(doc, "Methods", "Regression, Classification, NLP, Time Series, A/B Testing, Bayesian Inference")

    add_section_heading(doc, "Professional Experience", color=(0, 128, 128), border_bottom=True)
    add_entry(doc, "[Data Scientist / Senior DS]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Built [model type] predicting [outcome] with [X]% accuracy, driving $[X]M in [savings/revenue]")
    add_bullet(doc, "Designed A/B testing framework used across [X] product teams, improving experiment velocity by [X]%")
    add_bullet(doc, "Deployed ML models to production serving [X]M predictions/day via [REST API/batch pipeline]")

    add_entry(doc, "[Data Analyst / Junior DS]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Created automated reporting dashboards reducing manual reporting by [X] hours/week")
    add_bullet(doc, "Performed cohort analysis identifying [insight] that increased retention by [X]%")

    add_section_heading(doc, "Publications & Talks", color=(0, 128, 128), border_bottom=True)
    add_body(doc, "[Paper Title] — [Conference/Journal] — [Year]")
    add_body(doc, "[Talk Title] — [Event/Meetup] — [Year]")

    add_section_heading(doc, "Education", color=(0, 128, 128), border_bottom=True)
    add_entry(doc, "[MS/PhD] in [Statistics / Computer Science / Data Science]", "[University]", "[Year]")
    add_entry(doc, "[BS] in [Mathematics / CS / Engineering]", "[University]", "[Year]")

    return doc

TEMPLATES.append(("data-science-ats-resume.docx", gen_data_science))


# ---------------------------------------------------------------------------
# 11. Executive Leadership
# ---------------------------------------------------------------------------
def gen_executive_leadership():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 26, (30, 30, 30))
    add_contact(doc, "[C-Suite Title] | [Email] | [Phone] | [LinkedIn]", size=11, color=(80, 80, 80))
    add_separator(doc)

    add_section_heading(doc, "Executive Profile", color=(30, 30, 30), border_bottom=True, size=13)
    add_body(doc, "Visionary [C-level title] with [X]+ years of leadership across [industries]. Proven track record of scaling organizations from $[X]M to $[X]B in revenue. Board-level experience with expertise in [strategic area 1], [strategic area 2], and [strategic area 3]. Known for building high-performance cultures and driving sustainable growth.")

    add_section_heading(doc, "Areas of Expertise", color=(30, 30, 30), border_bottom=True, size=13)
    add_body(doc, "Strategic Planning • P&L Management • Digital Transformation • M&A Integration\nTalent Development • Board Relations • Investor Relations • Change Management")

    add_section_heading(doc, "Career Highlights", color=(30, 30, 30), border_bottom=True, size=13)
    add_bullet(doc, "Grew enterprise value from $[X]M to $[X]B through [organic growth/acquisitions]")
    add_bullet(doc, "Led IPO/fundraising of $[X]M at [valuation]")
    add_bullet(doc, "Transformed [division] from [loss-making] to [X]% profit margin in [X] years")
    add_bullet(doc, "Built and led global team of [X]+ across [X] countries")

    add_section_heading(doc, "Professional Experience", color=(30, 30, 30), border_bottom=True, size=13)
    add_entry(doc, "[C-Suite Title, e.g., CEO / CTO / COO / CFO]", "[Company Name] — [Location]", "[Start] – Present")
    add_bullet(doc, "Full P&L responsibility for $[X]B [division/company] with [X]+ employees")
    add_bullet(doc, "Developed and executed [X]-year strategic roadmap achieving [X]% YoY growth")
    add_bullet(doc, "Drove digital transformation initiative reducing operational costs by $[X]M")

    add_entry(doc, "[VP / SVP / EVP Title]", "[Company Name] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Managed [X] direct reports and [X]+ total team members")
    add_bullet(doc, "Secured [X] enterprise accounts generating $[X]M in annual recurring revenue")

    add_section_heading(doc, "Education", color=(30, 30, 30), border_bottom=True, size=13)
    add_entry(doc, "MBA", "[Top Business School]", "[Year]")
    add_entry(doc, "[Bachelor's] in [Field]", "[University]", "[Year]")

    add_section_heading(doc, "Board Memberships", color=(30, 30, 30), border_bottom=True, size=13)
    add_body(doc, "[Company/Organization] — [Board Member / Advisor] — [Year]–Present")

    return doc

TEMPLATES.append(("executive-leadership-ats-resume.docx", gen_executive_leadership))


# ---------------------------------------------------------------------------
# 12. Senior Manager
# ---------------------------------------------------------------------------
def gen_senior_manager():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (0, 82, 155))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [City, State]", color=(100, 100, 100))
    add_separator(doc)

    add_section_heading(doc, "Professional Summary", color=(0, 82, 155), border_bottom=True)
    add_body(doc, "Results-oriented [Director/Senior Manager] with [X]+ years of experience driving [function, e.g., operations, marketing, finance] excellence. Track record of leading [X]+ person teams, managing $[X]M budgets, and delivering [X]% improvement in [key metric]. Expert in [domain 1] and [domain 2].")

    add_section_heading(doc, "Core Competencies", color=(0, 82, 155), border_bottom=True)
    add_body(doc, "Team Leadership • Budget Management • Strategic Planning • Process Optimization\nStakeholder Management • Data-Driven Decision Making • Cross-Functional Collaboration • Vendor Management")

    add_section_heading(doc, "Professional Experience", color=(0, 82, 155), border_bottom=True)
    add_entry(doc, "[Director / Senior Manager]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Lead team of [X] across [X] functions delivering [products/services] to [X]+ clients")
    add_bullet(doc, "Increased [revenue/efficiency/NPS] by [X]% through [strategy/initiative]")
    add_bullet(doc, "Managed $[X]M annual budget with [X]% cost reduction through [optimization]")
    add_bullet(doc, "Implemented [system/process] that reduced [cycle time/errors] by [X]%")

    add_entry(doc, "[Manager / Assistant Director]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Promoted from [previous role] within [X] months based on [achievement]")
    add_bullet(doc, "Delivered [X] projects on time and within budget over [X]-year period")
    add_bullet(doc, "Established KPI framework adopted by [X] departments across the organization")

    add_entry(doc, "[Associate / Analyst / Coordinator]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Contributed to [major initiative] resulting in [measurable outcome]")
    add_bullet(doc, "Recognized with [award] for [performance/contribution]")

    add_section_heading(doc, "Education", color=(0, 82, 155), border_bottom=True)
    add_entry(doc, "[MBA / Master's] in [Field]", "[University]", "[Year]")
    add_entry(doc, "[Bachelor's] in [Field]", "[University]", "[Year]")

    add_section_heading(doc, "Certifications", color=(0, 82, 155), border_bottom=True)
    add_body(doc, "[PMP / Six Sigma / CPA / etc.] — [Issuing Body] — [Year]")

    return doc

TEMPLATES.append(("senior-manager-ats-resume.docx", gen_senior_manager))


# ---------------------------------------------------------------------------
# 13. Modern Two-Tone
# ---------------------------------------------------------------------------
def gen_modern_two_tone():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 24, (52, 73, 94))
    add_contact(doc, "[Target Role Title]", "Calibri", 12, color=(52, 73, 94))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [City, State]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "About Me", color=(52, 73, 94), border_bottom=True)
    add_body(doc, "Dynamic [profession] combining [X] years of [area 1] expertise with a passion for [area 2]. I thrive in fast-paced environments where I can [value proposition]. My approach blends [quality 1] with [quality 2] to deliver [type of results].")

    add_section_heading(doc, "Experience", color=(52, 73, 94), border_bottom=True)
    add_entry(doc, "[Job Title]", "[Company Name] — [Location]", "[Start] – Present")
    add_bullet(doc, "Spearheaded [initiative] resulting in [X]% growth in [metric]")
    add_bullet(doc, "Designed [system/process] adopted by [X] teams across the organization")
    add_bullet(doc, "Collaborated with [stakeholders] to launch [product/campaign] reaching [X]K+ users")

    add_entry(doc, "[Job Title]", "[Company Name] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Delivered [X] projects with combined value of $[X]M")
    add_bullet(doc, "Introduced [tool/practice] that improved team productivity by [X]%")

    add_section_heading(doc, "Education", color=(52, 73, 94), border_bottom=True)
    add_entry(doc, "[Degree] in [Major]", "[University]", "[Year]")

    add_section_heading(doc, "Skills", color=(52, 73, 94), border_bottom=True)
    add_body(doc, "[Skill 1] • [Skill 2] • [Skill 3] • [Skill 4] • [Skill 5] • [Skill 6] • [Skill 7] • [Skill 8] • [Skill 9] • [Skill 10]")

    add_section_heading(doc, "Languages", color=(52, 73, 94), border_bottom=True)
    add_body(doc, "[Language 1] (Native) • [Language 2] (Fluent) • [Language 3] (Conversational)")

    return doc

TEMPLATES.append(("modern-two-tone-ats-resume.docx", gen_modern_two_tone))


# ---------------------------------------------------------------------------
# 14. Simple Elegant
# ---------------------------------------------------------------------------
def gen_simple_elegant():
    doc = Document()
    set_margins(doc, 0.8, 0.8, 0.8, 0.8)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_contact(doc, "[City, State] • [Phone] • [Email] • [LinkedIn]", "Calibri", 10, WD_ALIGN_PARAGRAPH.CENTER, color=(130, 130, 130))
    add_separator(doc)

    add_section_heading(doc, "Profile", border_bottom=True, color=(80, 80, 80))
    add_body(doc, "Detail-oriented [profession] with [X] years of experience delivering [type of work] in [industry]. Adept at [skill 1] and [skill 2], with a commitment to [quality/value proposition].")

    add_section_heading(doc, "Experience", border_bottom=True, color=(80, 80, 80))
    add_entry(doc, "[Job Title]", "[Company], [Location]", "[Start] – [End]")
    add_bullet(doc, "Managed [responsibility] resulting in [X]% improvement in [metric]")
    add_bullet(doc, "Prepared [deliverables] for [audience], receiving [positive feedback/award]")
    add_bullet(doc, "Coordinated with [X] departments to ensure [outcome]")

    add_entry(doc, "[Job Title]", "[Company], [Location]", "[Start] – [End]")
    add_bullet(doc, "Supported [function] by [specific action], contributing to [result]")
    add_bullet(doc, "Maintained [system/records] with [X]% accuracy rate")

    add_section_heading(doc, "Education", border_bottom=True, color=(80, 80, 80))
    add_entry(doc, "[Degree] in [Major]", "[University], [Location]", "[Year]")

    add_section_heading(doc, "Skills", border_bottom=True, color=(80, 80, 80))
    add_body(doc, "[Skill 1], [Skill 2], [Skill 3], [Skill 4], [Skill 5], [Skill 6], [Skill 7], [Skill 8]")

    add_section_heading(doc, "Volunteer Work", border_bottom=True, color=(80, 80, 80))
    add_body(doc, "[Organization] — [Role] — [Date Range]")

    return doc

TEMPLATES.append(("simple-elegant-ats-resume.docx", gen_simple_elegant))


# ---------------------------------------------------------------------------
# 15. Chronological Professional
# ---------------------------------------------------------------------------
def gen_chronological_professional():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (0, 100, 0))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [City, State]", color=(100, 100, 100))
    add_separator(doc)

    add_section_heading(doc, "Professional Summary", color=(0, 100, 0), border_bottom=True)
    add_body(doc, "Seasoned [profession] with a [X]-year track record of progressive career growth in [industry]. Currently [current role] at [company type], responsible for [primary responsibility]. Consistently recognized for [distinguishing quality].")

    add_section_heading(doc, "Professional Experience", color=(0, 100, 0), border_bottom=True)

    add_entry(doc, "[Current Title]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Direct [scope of responsibility] generating $[X]M in annual revenue")
    add_bullet(doc, "Implemented [system/strategy] that increased [KPI] by [X]% YoY")
    add_bullet(doc, "Manage relationships with [X]+ key accounts/stakeholders")

    add_entry(doc, "[Previous Title]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Promoted to [role] after [achievement] within [timeframe]")
    add_bullet(doc, "Launched [initiative] contributing $[X]M in new business")

    add_entry(doc, "[Earlier Title]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Started as [entry role] and advanced through [X] positions over [X] years")
    add_bullet(doc, "Awarded [recognition] for [specific contribution]")

    add_entry(doc, "[Entry-Level Title]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Gained foundational skills in [area] supporting [function]")

    add_section_heading(doc, "Education", color=(0, 100, 0), border_bottom=True)
    add_entry(doc, "[Degree] in [Major]", "[University]", "[Year]")

    add_section_heading(doc, "Professional Development", color=(0, 100, 0), border_bottom=True)
    add_body(doc, "[Certification] — [Year] | [Course/Training] — [Year]")

    add_section_heading(doc, "Skills", color=(0, 100, 0), border_bottom=True)
    add_body(doc, "[Skill 1] • [Skill 2] • [Skill 3] • [Skill 4] • [Skill 5] • [Skill 6]")

    return doc

TEMPLATES.append(("chronological-professional-ats-resume.docx", gen_chronological_professional))


# ---------------------------------------------------------------------------
# 16. Career Changer
# ---------------------------------------------------------------------------
def gen_career_changer():
    doc = Document()
    set_margins(doc)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (153, 51, 0))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [City, State]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Career Objective", color=(153, 51, 0), border_bottom=True)
    add_body(doc, "[Current profession] transitioning to [target field] with [X] years of transferable experience in [relevant skill 1], [relevant skill 2], and [relevant skill 3]. Completed [certification/bootcamp/degree] in [new field]. Eager to bring [unique perspective from previous career] to [target role].")

    add_section_heading(doc, "Transferable Skills", color=(153, 51, 0), border_bottom=True)
    add_body(doc, "[Skill from old career relevant to new] • [Skill 2] • [Skill 3] • [Skill 4]\n[New skill from training] • [New skill 2] • [New skill 3] • [New skill 4]")

    add_section_heading(doc, "Relevant Training & Certifications", color=(153, 51, 0), border_bottom=True)
    add_entry(doc, "[Certification/Bootcamp/Degree in New Field]", "[Institution]", "[Year]")
    add_bullet(doc, "Completed [X]-hour program covering [topic 1], [topic 2], [topic 3]")
    add_bullet(doc, "Capstone project: [Project name] — [description] — [result/link]")

    add_section_heading(doc, "Relevant Projects", color=(153, 51, 0), border_bottom=True)
    add_entry(doc, "[Project Name]", "[Technologies/Tools Used]", "[Date]")
    add_bullet(doc, "[What you built/did] demonstrating [skill relevant to new career]")

    add_section_heading(doc, "Professional Experience", color=(153, 51, 0), border_bottom=True)
    add_entry(doc, "[Previous Career Title]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "[Transferable achievement, e.g., Managed $XM budget → financial acumen]")
    add_bullet(doc, "[Transferable achievement, e.g., Led team of X → leadership]")
    add_bullet(doc, "[Transferable achievement, e.g., Analyzed data to drive decisions → analytical skills]")

    add_entry(doc, "[Previous Career Title]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "[Highlight transferable skills with quantified results]")

    add_section_heading(doc, "Education", color=(153, 51, 0), border_bottom=True)
    add_entry(doc, "[Original Degree] in [Original Major]", "[University]", "[Year]")

    return doc

TEMPLATES.append(("career-changer-ats-resume.docx", gen_career_changer))


# ---------------------------------------------------------------------------
# 17. DevOps Engineer
# ---------------------------------------------------------------------------
def gen_devops_engineer():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (230, 126, 34))
    add_contact(doc, "[Email] | [Phone] | [GitHub] | [LinkedIn] | [Location]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Summary", color=(192, 100, 20), border_bottom=True)
    add_body(doc, "DevOps/SRE Engineer with [X]+ years of experience designing, automating, and optimizing cloud infrastructure. Expert in [cloud provider], CI/CD, and infrastructure-as-code. Reduced deployment frequency from [X/month] to [X/day] while maintaining [X]% uptime.")

    add_section_heading(doc, "Technical Skills", color=(192, 100, 20), border_bottom=True)
    add_skills_line(doc, "Cloud Platforms", "AWS (EC2, ECS, Lambda, S3, RDS, CloudFormation), GCP, Azure")
    add_skills_line(doc, "Containers & Orchestration", "Docker, Kubernetes, Helm, ECS, Docker Compose")
    add_skills_line(doc, "CI/CD", "GitHub Actions, Jenkins, GitLab CI, ArgoCD, CircleCI")
    add_skills_line(doc, "IaC & Config Mgmt", "Terraform, Ansible, Pulumi, CloudFormation")
    add_skills_line(doc, "Monitoring", "Prometheus, Grafana, Datadog, PagerDuty, ELK Stack")
    add_skills_line(doc, "Scripting", "Python, Bash, Go, PowerShell")
    add_skills_line(doc, "Networking", "VPC, DNS, Load Balancers, CDN, VPN, Firewalls")

    add_section_heading(doc, "Professional Experience", color=(192, 100, 20), border_bottom=True)
    add_entry(doc, "[Senior DevOps Engineer / SRE]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Architected multi-region [AWS/GCP] infrastructure serving [X]M+ daily active users")
    add_bullet(doc, "Reduced infrastructure costs by [X]% ($[X]K/month) through right-sizing and reserved instances")
    add_bullet(doc, "Built zero-downtime deployment pipeline supporting [X]+ microservices")
    add_bullet(doc, "Achieved [X]% uptime SLA by implementing [monitoring/alerting/auto-scaling strategy]")
    add_bullet(doc, "Automated [X]+ manual runbook procedures, saving [X] hours/week of toil")

    add_entry(doc, "[DevOps Engineer / Cloud Engineer]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Migrated [X] applications from on-premise to [cloud provider] in [X] months")
    add_bullet(doc, "Implemented Terraform modules for [X]+ infrastructure components with [X]% code reuse")
    add_bullet(doc, "Set up centralized logging and monitoring for [X]+ services using [ELK/Datadog]")

    add_section_heading(doc, "Certifications", color=(192, 100, 20), border_bottom=True)
    add_body(doc, "[AWS Solutions Architect / CKA / GCP Professional / etc.] — [Year]")

    add_section_heading(doc, "Education", color=(192, 100, 20), border_bottom=True)
    add_entry(doc, "[BS/MS] in [Computer Science / IT / Engineering]", "[University]", "[Year]")

    return doc

TEMPLATES.append(("devops-engineer-ats-resume.docx", gen_devops_engineer))


# ---------------------------------------------------------------------------
# 18. Product Manager
# ---------------------------------------------------------------------------
def gen_product_manager():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 22, (142, 68, 173))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn] | [Portfolio/Blog] | [City, State]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Summary", color=(142, 68, 173), border_bottom=True)
    add_body(doc, "Product Manager with [X]+ years of experience shipping [B2B/B2C/SaaS] products at scale. Track record of driving [X]% growth in [key metric] through data-informed product strategy. Experienced in [0-to-1 / growth / platform] product development. Skilled at translating customer insights into product roadmaps.")

    add_section_heading(doc, "Core Competencies", color=(142, 68, 173), border_bottom=True)
    add_body(doc, "Product Strategy • User Research • A/B Testing • PRDs & Specs • Agile/Scrum\nData Analysis • Stakeholder Management • Go-to-Market • OKRs • Wireframing")

    add_section_heading(doc, "Professional Experience", color=(142, 68, 173), border_bottom=True)
    add_entry(doc, "[Senior Product Manager / PM Lead]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Own product roadmap for [product/feature] with $[X]M ARR and [X]K+ MAU")
    add_bullet(doc, "Drove [X]% increase in [conversion/engagement/retention] through [feature/experiment]")
    add_bullet(doc, "Led cross-functional team of [X] engineers, [X] designers, and [X] data scientists")
    add_bullet(doc, "Defined and tracked OKRs achieving [X]% goal attainment rate over [X] quarters")
    add_bullet(doc, "Conducted [X]+ user interviews and synthesized insights into [X] product improvements")

    add_entry(doc, "[Product Manager / Associate PM]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Launched [feature/product] from 0 to [X]K users in [X] months")
    add_bullet(doc, "Reduced churn by [X]% through [customer feedback loop / feature improvement]")
    add_bullet(doc, "Wrote PRDs, user stories, and acceptance criteria for [X]+ features per quarter")

    add_section_heading(doc, "Education", color=(142, 68, 173), border_bottom=True)
    add_entry(doc, "[MBA / MS] in [Business / CS / HCI]", "[University]", "[Year]")
    add_entry(doc, "[BS] in [Field]", "[University]", "[Year]")

    add_section_heading(doc, "Tools", color=(142, 68, 173), border_bottom=True)
    add_body(doc, "Jira, Figma, Amplitude, Mixpanel, SQL, Notion, Miro, Linear, Productboard")

    return doc

TEMPLATES.append(("product-manager-ats-resume.docx", gen_product_manager))


# ---------------------------------------------------------------------------
# 19. Executive VP
# ---------------------------------------------------------------------------
def gen_executive_vp():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 26, (44, 62, 80))
    add_contact(doc, "[VP / SVP / EVP Title]", "Calibri", 13, color=(44, 62, 80))
    add_contact(doc, "[Email] | [Phone] | [LinkedIn]", color=(100, 100, 100))
    add_separator(doc)

    add_section_heading(doc, "Executive Summary", color=(44, 62, 80), border_bottom=True, size=13)
    add_body(doc, "Transformational [VP-level title] with [X]+ years leading [function] at [Fortune 500 / high-growth / private equity-backed] companies. Expert in [domain 1] and [domain 2]. Track record of building teams, scaling operations, and delivering [X]% improvement in [business metric].")

    add_section_heading(doc, "Key Accomplishments", color=(44, 62, 80), border_bottom=True, size=13)
    add_bullet(doc, "Scaled [function/team] from [X] to [X]+ FTEs across [X] global offices")
    add_bullet(doc, "Delivered $[X]M in annual cost savings through [operational excellence initiative]")
    add_bullet(doc, "Drove [product/market] expansion into [X] new markets generating $[X]M revenue")
    add_bullet(doc, "Led post-acquisition integration of [X] companies within [X] months")

    add_section_heading(doc, "Professional Experience", color=(44, 62, 80), border_bottom=True, size=13)
    add_entry(doc, "[VP / SVP / EVP of Function]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Report to [CEO/President/Board] with oversight of [X] departments and $[X]M budget")
    add_bullet(doc, "Designed [strategy/program] that increased [metric] by [X]% in first [X] months")
    add_bullet(doc, "Championed [initiative] recognized as [industry best practice/award]")

    add_entry(doc, "[VP / Director]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Built [function/team] from scratch, hiring [X] professionals in [X] months")
    add_bullet(doc, "Achieved [X]% year-over-year revenue growth for [X] consecutive years")

    add_section_heading(doc, "Education", color=(44, 62, 80), border_bottom=True, size=13)
    add_entry(doc, "MBA", "[Top Business School]", "[Year]")
    add_entry(doc, "[Bachelor's] in [Field]", "[University]", "[Year]")

    add_section_heading(doc, "Board & Industry Involvement", color=(44, 62, 80), border_bottom=True, size=13)
    add_body(doc, "[Organization] — [Board Member / Committee Chair] — [Year]–Present")
    add_body(doc, "[Industry Association] — [Officer / Speaker] — [Year]–Present")

    return doc

TEMPLATES.append(("executive-vp-ats-resume.docx", gen_executive_vp))


# ---------------------------------------------------------------------------
# 20. Creative Professional
# ---------------------------------------------------------------------------
def gen_creative_professional():
    doc = Document()
    set_margins(doc, 0.6, 0.6, 0.7, 0.7)
    add_name(doc, "[YOUR FULL NAME]", "Calibri", 24, (192, 57, 43))
    add_contact(doc, "[Creative Title, e.g., UX Designer / Content Strategist / Art Director]", "Calibri", 12, color=(192, 57, 43))
    add_contact(doc, "[Email] | [Portfolio URL] | [LinkedIn] | [Phone]", color=(120, 120, 120))
    add_separator(doc)

    add_section_heading(doc, "Creative Profile", color=(192, 57, 43), border_bottom=True)
    add_body(doc, "Award-winning [creative role] with [X]+ years of experience in [industry, e.g., tech, media, advertising]. Passionate about [design philosophy / creative approach]. Portfolio includes work for [notable brands/companies]. Skilled in translating business objectives into compelling [visual/written/interactive] experiences.")

    add_section_heading(doc, "Core Skills", color=(192, 57, 43), border_bottom=True)
    add_skills_line(doc, "Design", "[Figma], [Sketch], [Adobe Creative Suite], [Framer], [Webflow]")
    add_skills_line(doc, "Technical", "[HTML/CSS], [JavaScript], [React], [After Effects], [Blender]")
    add_skills_line(doc, "Strategy", "[User Research], [A/B Testing], [Brand Strategy], [Content Strategy]")
    add_skills_line(doc, "Methods", "[Design Thinking], [Design Systems], [Accessibility (WCAG)], [Agile]")

    add_section_heading(doc, "Professional Experience", color=(192, 57, 43), border_bottom=True)
    add_entry(doc, "[Senior UX Designer / Creative Director / Lead Designer]", "[Company] — [Location]", "[Start] – Present")
    add_bullet(doc, "Led design for [product/brand] used by [X]M+ users, increasing [engagement/conversion] by [X]%")
    add_bullet(doc, "Created and maintained design system with [X]+ components adopted by [X] product teams")
    add_bullet(doc, "Conducted [X]+ user research sessions informing [X] major feature launches")
    add_bullet(doc, "Won [award/recognition] for [project/campaign]")

    add_entry(doc, "[UX Designer / Graphic Designer / Content Creator]", "[Company] — [Location]", "[Start] – [End]")
    add_bullet(doc, "Designed [X]+ [screens/pages/campaigns] for [product/client portfolio]")
    add_bullet(doc, "Improved [accessibility/load time/usability score] by [X]% through [action]")
    add_bullet(doc, "Collaborated with [engineers/copywriters/marketers] to deliver [project] on deadline")

    add_section_heading(doc, "Selected Projects", color=(192, 57, 43), border_bottom=True)
    add_entry(doc, "[Project/Brand Name]", "[Your Role]", "[Link to portfolio piece]")
    add_body(doc, "[Brief description of project, your contribution, and the impact/result]")

    add_section_heading(doc, "Education", color=(192, 57, 43), border_bottom=True)
    add_entry(doc, "[BFA/MFA/BS] in [Design / Fine Arts / HCI / Communications]", "[University]", "[Year]")

    add_section_heading(doc, "Awards & Recognition", color=(192, 57, 43), border_bottom=True)
    add_body(doc, "[Award Name] — [Organization] — [Year]")
    add_body(doc, "[Award Name] — [Organization] — [Year]")

    return doc

TEMPLATES.append(("creative-professional-ats-resume.docx", gen_creative_professional))


# =============================================================================
# GENERATE ALL
# =============================================================================
if __name__ == "__main__":
    print(f"Generating {len(TEMPLATES)} DOCX templates...\n")
    for filename, gen_fn in TEMPLATES:
        doc = gen_fn()
        path = os.path.join(OUT_DIR, filename)
        doc.save(path)
        size_kb = os.path.getsize(path) / 1024
        print(f"  {filename:50s} ({size_kb:.1f} KB)")
    print(f"\nDone! {len(TEMPLATES)} templates saved to {OUT_DIR}")
